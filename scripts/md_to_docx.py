"""
Convert Markdown files to Word (.docx) for MediVoice VN project.
Usage: python scripts/md_to_docx.py
Output: exports/ folder
"""

import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = r"D:\MediVoice_VN\exports"
FILES = [
    (r"D:\MediVoice_VN\docs\cl08_operation\BRS.md", "BRS_MediVoice_VN.docx"),
    (r"D:\MediVoice_VN\docs\cl05_leadership\VISION.md", "VISION_MediVoice_VN.docx"),
    (r"D:\MediVoice_VN\docs\records\THIRD_PARTY_REVIEW_REQUEST.md", "THIRD_PARTY_REVIEW_REQUEST.docx"),
]


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_bold_run(para, text):
    run = para.add_run(text)
    run.bold = True
    return run


def parse_inline(para, text):
    """Parse **bold** and `code` inline markup."""
    pattern = re.compile(r"(\*\*(.+?)\*\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(2):  # bold
            r = para.add_run(m.group(2))
            r.bold = True
        elif m.group(3):  # code
            r = para.add_run(m.group(3))
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def convert_md_to_docx(md_path, out_path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    i = 0

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        # Filter separator rows
        data_rows = [r for r in table_rows if not re.match(r"^\|[-| :]+\|$", r.strip())]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            parsed.append(cells)
        max_cols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=max_cols)
        tbl.style = "Table Grid"
        for ri, row in enumerate(parsed):
            for ci, cell_text in enumerate(row):
                if ci >= max_cols:
                    break
                cell = tbl.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                parse_inline(p, cell_text)
                if ri == 0:
                    set_cell_bg(cell, "D9E1F2")
                    for run in p.runs:
                        run.bold = True
                for run in p.runs:
                    run.font.size = Pt(11)
        doc.add_paragraph()
        table_rows = []
        in_table = False

    def flush_code(lines_buf):
        if not lines_buf:
            return
        para = doc.add_paragraph()
        para.style = doc.styles["Normal"]
        run = para.add_run("\n".join(lines_buf))
        run.font.name = "Courier New"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            bdr = OxmlElement(f"w:{side}")
            bdr.set(qn("w:val"), "single")
            bdr.set(qn("w:sz"), "4")
            bdr.set(qn("w:space"), "4")
            bdr.set(qn("w:color"), "AAAAAA")
            pBdr.append(bdr)
        pPr.append(pBdr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F5F5F5")
        pPr.append(shd)

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Code block toggle
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                flush_code(code_lines)
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|"):
            if in_table:
                table_rows.append(line)
            else:
                in_table = True
                table_rows = [line]
            i += 1
            # Peek ahead for more table rows
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i].rstrip("\n"))
                i += 1
            flush_table()
            continue
        else:
            if in_table:
                flush_table()

        stripped = line.strip()

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "4472C4")
            pBdr.append(bot)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        if stripped.startswith("# ") and not stripped.startswith("## "):
            # Skip metadata comment headers (lines 1-3)
            text = stripped[2:].strip()
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            i += 1
            continue

        if stripped.startswith("## "):
            text = stripped[3:].strip()
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = RGBColor(0x26, 0x72, 0xC4)
            i += 1
            continue

        if stripped.startswith("### "):
            text = stripped[4:].strip()
            p = doc.add_heading(text, level=3)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph()
            p.style = doc.styles["Normal"]
            r = p.add_run(text)
            r.italic = True
            r.font.size = Pt(13)
            r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "720")
            pPr.append(ind)
            i += 1
            continue

        # Bullet list
        if stripped.startswith("- "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, text)
            for run in p.runs:
                run.font.size = Pt(12)
            i += 1
            continue

        # Empty line
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        parse_inline(p, stripped)
        i += 1

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(out_path)
    print(f"  Saved: {out_path}")


if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    for md_file, docx_name in FILES:
        out = os.path.join(OUTPUT_DIR, docx_name)
        print(f"Converting: {os.path.basename(md_file)} -> {docx_name}")
        convert_md_to_docx(md_file, out)
    print("\nDone. Files saved to D:\\MediVoice_VN\\exports\\")
